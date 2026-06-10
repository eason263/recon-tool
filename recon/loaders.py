"""Load files into one of two normalized representations.

- "lines": list[str] for text-like files (structured formats canonicalized first)
- "records": (headers: list[str], rows: list[dict]) for tabular files
"""

import csv
from pathlib import Path

from .formatters import pretty_json, pretty_xml


class LoadError(Exception):
    pass


def load_text_lines(path: Path) -> list[str]:
    return _read_text(path).splitlines()


def load_json_lines(path: Path) -> list[str]:
    try:
        return pretty_json(_read_text(path)).splitlines()
    except ValueError as e:
        raise LoadError(f"{path}: invalid JSON: {e}") from e


def load_xml_lines(path: Path) -> list[str]:
    try:
        return pretty_xml(_read_text(path)).splitlines()
    except Exception as e:
        raise LoadError(f"{path}: invalid XML: {e}") from e


def load_docx_lines(path: Path) -> list[str]:
    import docx

    try:
        document = docx.Document(str(path))
    except Exception as e:
        raise LoadError(f"{path}: cannot read .docx: {e}") from e
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return lines


def load_csv_records(path: Path) -> tuple[list[str], list[dict]]:
    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        if reader.fieldnames is None:
            raise LoadError(f"{path}: empty CSV, no header row")
        rows = [dict(row) for row in reader]
    return list(reader.fieldnames), rows


def load_xlsx_records(path: Path) -> tuple[list[str], list[dict]]:
    import openpyxl

    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as e:
        raise LoadError(f"{path}: cannot read .xlsx: {e}") from e
    ws = wb.active
    rows_iter = ws.iter_rows(values_only=True)
    try:
        header_row = next(rows_iter)
    except StopIteration:
        raise LoadError(f"{path}: empty worksheet, no header row") from None
    headers = [str(h) if h is not None else f"col{i+1}" for i, h in enumerate(header_row)]
    rows = []
    for raw in rows_iter:
        if raw is None or all(v is None for v in raw):
            continue
        rows.append({h: ("" if v is None else str(v)) for h, v in zip(headers, raw)})
    wb.close()
    return headers, rows


# extension -> (kind, loader). kind is "lines" or "records".
LOADERS = {
    ".txt": ("lines", load_text_lines),
    ".log": ("lines", load_text_lines),
    ".java": ("lines", load_text_lines),
    ".json": ("lines", load_json_lines),
    ".xml": ("lines", load_xml_lines),
    ".docx": ("lines", load_docx_lines),
    ".csv": ("records", load_csv_records),
    ".xlsx": ("records", load_xlsx_records),
}


def load(path: str | Path) -> tuple[str, object]:
    """Return (kind, data) for a file, falling back to plain text for unknown extensions."""
    path = Path(path)
    if not path.is_file():
        raise LoadError(f"{path}: file not found")
    ext = path.suffix.lower()
    kind, loader = LOADERS.get(ext, ("lines", load_text_lines))
    return kind, loader(path)


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="latin-1")
        except Exception as e:
            raise LoadError(f"{path}: not a readable text file: {e}") from e
