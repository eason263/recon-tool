from pathlib import Path

from recon.loaders import load

FIXTURES = Path(__file__).parent / "fixtures"


def test_json_canonicalization_ignores_key_order_and_whitespace(tmp_path):
    a = tmp_path / "a.json"
    b = tmp_path / "b.json"
    a.write_text('{"b": 2, "a": 1}')
    b.write_text('{\n  "a": 1,\n  "b": 2\n}')
    assert load(a) == load(b)
    kind, lines = load(a)
    assert kind == "lines"
    assert lines == ['{', '  "a": 1,', '  "b": 2', '}']


def test_xml_canonicalization_ignores_whitespace(tmp_path):
    a = tmp_path / "a.xml"
    b = tmp_path / "b.xml"
    a.write_text("<root><item>1</item><item>2</item></root>")
    b.write_text("<root>\n   <item>1</item>\n   <item>2</item>\n</root>")
    assert load(a) == load(b)
    kind, lines = load(a)
    assert kind == "lines"
    assert any("<item>1</item>" in line for line in lines)


def test_csv_records(tmp_path):
    f = tmp_path / "x.csv"
    f.write_text("id,name\n1,alice\n2,bob\n")
    kind, (headers, rows) = load(f)
    assert kind == "records"
    assert headers == ["id", "name"]
    assert rows == [{"id": "1", "name": "alice"}, {"id": "2", "name": "bob"}]


def test_xlsx_records(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["id", "amount"])
    ws.append([1, 10.5])
    ws.append([2, None])
    f = tmp_path / "x.xlsx"
    wb.save(f)

    kind, (headers, rows) = load(f)
    assert kind == "records"
    assert headers == ["id", "amount"]
    assert rows == [{"id": "1", "amount": "10.5"}, {"id": "2", "amount": ""}]


def test_docx_lines(tmp_path):
    import docx

    document = docx.Document()
    document.add_paragraph("Hello")
    document.add_paragraph("World")
    f = tmp_path / "x.docx"
    document.save(f)

    kind, lines = load(f)
    assert kind == "lines"
    assert "Hello" in lines and "World" in lines


def test_unknown_extension_falls_back_to_text(tmp_path):
    f = tmp_path / "x.weird"
    f.write_text("line1\nline2\n")
    assert load(f) == ("lines", ["line1", "line2"])
