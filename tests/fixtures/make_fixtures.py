"""Regenerate the binary fixtures (xlsx, docx). Run from this directory."""

import docx
import openpyxl

wb = openpyxl.Workbook()
ws = wb.active
ws.append(["trade_id", "ccy", "amount", "updated_at"])
ws.append(["T001", "USD", 1500.00, "2026-06-01"])
ws.append(["T002", "EUR", 230.50, "2026-06-01"])
ws.append(["T003", "GBP", 99.99, "2026-06-02"])
wb.save("a.xlsx")

wb = openpyxl.Workbook()
ws = wb.active
ws.append(["trade_id", "ccy", "amount", "updated_at"])
ws.append(["T001", "USD", 1500.005, "2026-06-03"])
ws.append(["T002", "EUR", 230.50, "2026-06-03"])
ws.append(["T004", "JPY", 12000, "2026-06-03"])
wb.save("b.xlsx")

doc = docx.Document()
doc.add_paragraph("Quarterly Report")
doc.add_paragraph("Revenue increased by 10%.")
doc.add_paragraph("Headcount is stable.")
doc.save("a.docx")

doc = docx.Document()
doc.add_paragraph("Quarterly Report")
doc.add_paragraph("Revenue increased by 12%.")
doc.add_paragraph("Headcount is stable.")
doc.add_paragraph("New office opened in Singapore.")
doc.save("b.docx")
